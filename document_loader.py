import os
import glob
from typing import List, Dict

from config import DATA_DIR, SUPPORTED_EXTENSIONS


def read_text_file(file_path: str) -> str:
    """
    读取 txt / md 文件。
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="gbk", errors="ignore") as f:
            return f.read()


def read_docx_file(file_path: str) -> str:
    """
    读取 docx 文件。
    主要读取正文段落和简单表格。
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("请先安装 python-docx：pip install python-docx") from exc

    doc = Document(file_path)

    paragraphs = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    for table_idx, table in enumerate(doc.tables):
        paragraphs.append(f"\n[Table {table_idx + 1}]")

        table_rows = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                row_text.append(cell_text)
            table_rows.append(row_text)

        markdown_table = table_to_markdown(table_rows)
        if markdown_table:
            paragraphs.append(markdown_table)

    return "\n\n".join(paragraphs)


def table_to_markdown(table):
    """
    将表格二维列表转换为 Markdown 表格。
    这样做的原因是：
    1. 保留表格行列关系；
    2. 让表格内容可以作为普通文本进入 RAG；
    3. 提高表格类问题的检索和回答效果。
    """
    if not table:
        return ""

    cleaned_rows = []

    for row in table:
        if row is None:
            continue

        cleaned_row = []
        for cell in row:
            if cell is None:
                cleaned_row.append("")
            else:
                cleaned_row.append(str(cell).replace("\n", " ").strip())

        if any(cell for cell in cleaned_row):
            cleaned_rows.append(cleaned_row)

    if not cleaned_rows:
        return ""

    # 统一列数，避免 Markdown 表格错位
    max_cols = max(len(row) for row in cleaned_rows)

    normalized_rows = []
    for row in cleaned_rows:
        if len(row) < max_cols:
            row = row + [""] * (max_cols - len(row))
        elif len(row) > max_cols:
            row = row[:max_cols]
        normalized_rows.append(row)

    header = normalized_rows[0]
    body = normalized_rows[1:]

    markdown = []
    markdown.append("| " + " | ".join(header) + " |")
    markdown.append("| " + " | ".join(["---"] * len(header)) + " |")

    for row in body:
        markdown.append("| " + " | ".join(row) + " |")

    return "\n".join(markdown)


def read_pdf_file(file_path: str) -> str:
    """
    读取文本型 PDF，并尝试提取表格。
    """
    try:
        import pdfplumber
    except ImportError as exc:
        raise ImportError("请先安装 pdfplumber：pip install pdfplumber") from exc

    all_pages = []

    with pdfplumber.open(file_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            page_parts = []
            page_no = page_idx + 1

            page_parts.append(f"[Page {page_no}]")

            # 1. 提取页面普通文本
            text = page.extract_text()
            if text:
                page_parts.append(text)

            # 2. 提取页面表格
            tables = page.extract_tables()

            for table_idx, table in enumerate(tables):
                markdown_table = table_to_markdown(table)

                if markdown_table:
                    page_parts.append(f"\n[Table {table_idx + 1} on Page {page_no}]")
                    page_parts.append(markdown_table)

            all_pages.append("\n\n".join(page_parts))

    return "\n\n".join(all_pages)


def read_document(file_path: str) -> str:
    """
    根据文件后缀调用不同读取函数。
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext in [".txt", ".md"]:
        return read_text_file(file_path)

    if ext == ".docx":
        return read_docx_file(file_path)

    if ext == ".pdf":
        return read_pdf_file(file_path)

    raise ValueError(f"暂不支持该文件类型：{ext}")


def load_documents(data_dir: str = DATA_DIR) -> List[Dict]:
    """
    读取 data 文件夹下所有支持格式文档。
    """
    documents = []

    all_files = []
    for ext in SUPPORTED_EXTENSIONS:
        all_files.extend(glob.glob(os.path.join(data_dir, f"*{ext}")))

    all_files = sorted(all_files)

    for file_path in all_files:
        text = read_document(file_path)

        if not text or not text.strip():
            print(f"[跳过] 文件内容为空或无法解析：{file_path}")
            continue

        documents.append({
            "source": os.path.basename(file_path),
            "path": file_path,
            "extension": os.path.splitext(file_path)[1].lower(),
            "text": text
        })

    return documents


if __name__ == "__main__":
    docs = load_documents()
    print(f"读取到文档数量：{len(docs)}")

    for doc in docs:
        print("=" * 80)
        print("source:", doc["source"])
        print("extension:", doc["extension"])
        print("text preview:")
        print(doc["text"][:1000])